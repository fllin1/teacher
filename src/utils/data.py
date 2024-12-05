"""
This module contains the classes and functions to handle data operations.
Current formats supported: JSON and docx.

Classes:
    Read
    Save
    File

Functions:
    Read.read_json(path) -> dict
    Read.read_docx(file_path) -> str
    Save.save_json(file, path) -> None
    File.get_files_name(folder_path: str = "../data/raw") -> list
"""

import json
import os

from docx import Document


class Read:
    """
    Class to read data.
    """

    @staticmethod
    def read_json(path) -> dict:
        """
        Read data from a JSON file.

        Args:
            path (Path): The path to the JSON file.

        Returns:
            dict: The loaded data.
        """
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
        print(f"Data loaded from {path}")

    @staticmethod
    def read_docx(file_path) -> str:
        """
        Read the docx file and return the full text

        Args:
            file_path (str): The file path of the docx file

        Returns:
            str: The full text of the docx file
        """
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return " ".join(full_text)


class Save:
    """
    Class to save data.
    """

    @staticmethod
    def save_json(file, path) -> None:
        """
        Save progress to a JSON file.

        Args:
            file (dict): The dictionary to save.
            path (Path): The path to the JSON file.
        """
        with open(path, "w", encoding="utf-8") as f:
            json.dump(file, f, indent=4, ensure_ascii=False)
        print(f"Progress saved to {path}")

    @staticmethod
    def save_docx(file, path) -> None:
        """
        Save progress to a docx file.

        Args:
            file (str): The string to save.
            path (Path): The path to the docx file.
        """
        doc = Document()
        doc.add_paragraph(file)
        doc.save(path)
        print(f"Progress saved to {path}")


class File:
    """
    Class to handle file operations.
    """

    @staticmethod
    def file_name(extension: str, folder_path: str = "../data/raw") -> list:
        """
        Get all the files in the folder_path

        Args:
            folder_path (str, optional): The folder path. Defaults to "../data/raw".

        Returns:
            list: List of files in the folder_path
        """
        return [
            f
            for f in os.listdir(folder_path)
            if os.path.isfile(os.path.join(folder_path, f))
            and f.endswith(f".{extension}")
        ]
